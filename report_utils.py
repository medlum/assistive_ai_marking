import zipfile
from pathlib import Path
import shutil
import re
import pandas as pd
import streamlit as st
from docx import Document
from pypdf import PdfReader
import io

def extract_and_read_files(zip_path):
    # Define extraction path
    extract_folder = "extracted_files"
    #extract_folder = st.session_state.user_id

    if Path(extract_folder).exists():
        shutil.rmtree(extract_folder)

    # Extract ZIP file
    with zipfile.ZipFile(zip_path, "r") as zip_ref:
            zip_ref.extractall(extract_folder)

    extracted_data = {}

    for folder in Path(extract_folder).iterdir(): 

        if folder.is_dir():  

        # extract student name from the each subfolder
        # which contains name as standard label from brightspace 
            folder_name = str(folder.relative_to(extract_folder))
            cleaned_text = re.sub(r"\b(BA|NP|PM|AM)\b", "", folder_name)
            cleaned_text = re.sub(r"\s+", " ", cleaned_text).strip()
            student_name = " ".join(re.findall(r"\b[A-Z]+\b", cleaned_text))

            
            for file in folder.glob("*.*"):

                if file.suffix.lower() == ".docx":
                    #doc = Document(file)
                    #data = "\n".join([para.text for para in doc.paragraphs])
                    #data = [[cell.text for cell in row.cells] for table in doc.tables for row in table.rows]

                    def ignore_images(image):
                        return {}
                    
                    doc = Document(file)
                    data = "\n".join([para.text for para in doc.paragraphs])
                    #result = mammoth.convert_to_html(file,convert_image=ignore_images)
                    #data = result.value
                
                elif file.suffix.lower() == ".pdf":
                    data = ""
                    reader = PdfReader(file)
                    for page in reader.pages:
                        data += page.extract_text()

                else:
                    st.error(".docx or .pdf files not found")
            
                if student_name not in extracted_data:
                    # store values as a list with file extension and extracted data
                    extracted_data[student_name] = [file.suffix.lower(), data]
    
    return extracted_data


    


#def extract_name_id(data):
#    pattern = r"Name\s*/\s*Student ID\s*\n(?:\s*\n)*([A-Za-z \-']+\s*/?\s*[A-Za-z0-9]+)"
#    match = re.search(pattern, data)
#    if match:
#        return match.group(1).strip()
#    else:
#        return "Name and student ID not found."
#
#def extract_and_read_files(uploaded_zip_file):
#    extracted_data = {}
#    data = ""
#
#    extract_folder = "extracted_files"
#    if Path(extract_folder).exists():
#        shutil.rmtree(extract_folder)
#
#    # ✅ Open from uploaded file (not file path)
#    with zipfile.ZipFile(io.BytesIO(uploaded_zip_file.read()), "r") as zip_ref:
#        zip_ref.extractall(extract_folder)
#
#    # ✅ Recursively process extracted files
#    for file in Path(extract_folder).rglob("*"):
#        if file.suffix.lower() == ".docx":
#            doc = Document(file)
#            data = "\n".join([para.text for para in doc.paragraphs])
#
#        elif file.suffix.lower() == ".pdf":
#            reader = PdfReader(file)
#            for page in reader.pages:
#                data += page.extract_text()
#
#        else:
#            continue  # skip unsupported files
#
#        student_name_id = extract_name_id(data)
#
#        if student_name_id not in extracted_data:
#            extracted_data[student_name_id] = [file.suffix.lower(), data]
#
#    return extracted_data



def process_data(data):
    # Convert list of dictionaries to DataFrame
    df = pd.DataFrame(data)

    df['Total'] = df.select_dtypes(include='number').sum(axis=1)

    return df

system_message = """
1. Your primary task is to evaluate students' written assignments based on a structured marking rubric.  
2. Follow the instructions to mark:
    - Refer closely to the provided marking rubric to ensure accurate and consistent grading.
    - Evaluate each criterion individually, assigning marks strictly according to the rubric.
    - Maintain a high academic standard throughout the assessment.
    - Do not exceed the maximum marks allocated for any criterion.
    - Provide detailed and constructive feedback, identifying specific strengths and weaknesses of the report.
    - Offer actionable suggestions for improvement in areas where the report falls short.
    - Highlight any sentences or sections that require revision or enhancement.
    - Justify the marks awarded for each criterion with clear, evidence-based reasoning.
    - Incorporate explanations of the mark allocation within the feedback for each criterion.
  
    - Return the marks and feedback in a dictionary : 
      {
          'Student Name': str,
          'Content (30 marks)': float,
          'Organisation and Structure (10 marks)': float,
          'Referencing (10 marks)': float,
          'Feedback' : str
      }  
    - Use single quotation '' for strings in the dictionary.
    - Your answer should only contain the returned dictionary and nothing else. 
"""

# custom CSS for buttons
btn_css = """
<style>
    .stButton > button {
        color: #383736; 
        border: none; /* No border */
        padding: 5px 22px; /* Reduced top and bottom padding */
        text-align: center; /* Centered text */
        text-decoration: none; /* No underline */
        display: inline-block; /* Inline-block */
        font-size: 8px !important;
        margin: 4px 2px; /* Some margin */
        cursor: pointer; /* Pointer cursor on hover */
        border-radius: 30px; /* Rounded corners */
        transition: background-color 0.3s; /* Smooth background transition */
    }
    .stButton > button:hover {
        color: #383736; 
        background-color: #c4c2c0; /* Darker green on hover */
    }
</style>
"""

image_css = """
<style>
.stImage img {
    border-radius: 50%;
    #border: 5px solid #f8fae6;
}
</style>

"""
